"""
Enhanced Dual RAG System with User-Specific Memory and Company Documents
"""

import os
import uuid
from typing import List, Dict, Any, Optional
from datetime import datetime
import json
from langchain_community.vectorstores import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import DirectoryLoader, TextLoader
from langchain_groq import ChatGroq
from langchain_core.documents import Document
from dotenv import load_dotenv

# Set protobuf implementation before any imports that might use it
if not os.getenv("PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION"):
    os.environ["PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION"] = "python"

# Import for Word document processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[ENHANCED_RAG] Warning: python-docx not available. Word document processing disabled.")

# Load environment variables from .env or config.env files
env_files = ['.env', 'config.env']
env_loaded = False

for env_file in env_files:
    if os.path.exists(env_file):
        load_dotenv(env_file)
        print(f"🔧 Enhanced RAG loaded environment variables from: {env_file}")
        env_loaded = True
        break

if not env_loaded:
    print("⚠️  Enhanced RAG: No .env or config.env file found, using system environment variables")

class UserChatHistoryRAG:
    """RAG system for user-specific chat history with namespacing"""
    
    def __init__(self, base_persist_directory: str = "user_chat_vectors"):
        self.base_persist_directory = base_persist_directory
        self.gemini_api_key = os.getenv("GEMINI_API_KEY")
        self.user_vectorstores = {}  # user_id -> vectorstore
        
        if self.gemini_api_key:
            self.embeddings = GoogleGenerativeAIEmbeddings(
                model="models/embedding-001",
                google_api_key=self.gemini_api_key
            )
        else:
            self.embeddings = None
            print("[USER_RAG] Warning: GEMINI_API_KEY not found. Chat history RAG disabled.")
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50
        )
        
        # Ensure base directory exists
        os.makedirs(base_persist_directory, exist_ok=True)
    
    def get_user_vectorstore(self, user_id: str):
        """Get or create vectorstore for a specific user"""
        if not self.embeddings:
            return None
            
        if user_id not in self.user_vectorstores:
            user_persist_dir = os.path.join(self.base_persist_directory, f"user_{user_id}")
            os.makedirs(user_persist_dir, exist_ok=True)
            
            try:
                self.user_vectorstores[user_id] = Chroma(
                    persist_directory=user_persist_dir,
                    embedding_function=self.embeddings,
                    collection_name=f"user_chat_{user_id}"
                )
                print(f"[USER_RAG] Created/loaded vectorstore for user {user_id}")
            except Exception as e:
                print(f"[USER_RAG] Error creating vectorstore for user {user_id}: {e}")
                return None
        
        return self.user_vectorstores[user_id]
    
    def store_conversation(self, user_id: str, user_message: str, agent_response: str, 
                          stage: str = "unknown", metadata: Dict = None):
        """Store a conversation turn in user's vector memory"""
        vectorstore = self.get_user_vectorstore(user_id)
        if not vectorstore:
            return
        
        timestamp = datetime.now().isoformat()
        conversation_text = f"User: {user_message}\nAgent: {agent_response}"
        
        doc_metadata = {
            "user_id": user_id,
            "timestamp": timestamp,
            "stage": stage,
            "user_message": user_message,
            "agent_response": agent_response,
            "conversation_id": str(uuid.uuid4()),
            **(metadata or {})
        }
        
        # Create document for this conversation turn
        doc = Document(
            page_content=conversation_text,
            metadata=doc_metadata
        )
        
        try:
            vectorstore.add_documents([doc])
            print(f"[USER_RAG] Stored conversation for user {user_id} at stage {stage}")
        except Exception as e:
            print(f"[USER_RAG] Error storing conversation for user {user_id}: {e}")
    
    def get_relevant_history(self, user_id: str, current_query: str, k: int = 5) -> List[Dict]:
        """Retrieve relevant conversation history for the user"""
        vectorstore = self.get_user_vectorstore(user_id)
        if not vectorstore:
            return []
        
        try:
            # Search for relevant conversations
            docs = vectorstore.similarity_search(current_query, k=k)
            
            history = []
            for doc in docs:
                history.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "timestamp": doc.metadata.get("timestamp"),
                    "stage": doc.metadata.get("stage"),
                    "user_message": doc.metadata.get("user_message"),
                    "agent_response": doc.metadata.get("agent_response")
                })
            
            # Sort by timestamp (most recent first)
            history.sort(key=lambda x: x.get("timestamp", ""), reverse=True)
            return history
            
        except Exception as e:
            print(f"[USER_RAG] Error retrieving history for user {user_id}: {e}")
            return []
    
    def get_user_conversation_summary(self, user_id: str, last_n: int = 10) -> str:
        """Get a summary of user's recent conversations"""
        vectorstore = self.get_user_vectorstore(user_id)
        if not vectorstore:
            return ""
        
        try:
            # Get recent conversations
            docs = vectorstore.similarity_search("conversation", k=last_n)
            if not docs:
                return ""
            
            # Sort by timestamp
            sorted_docs = sorted(docs, key=lambda x: x.metadata.get("timestamp", ""))
            
            summary_parts = []
            for doc in sorted_docs[-last_n:]:  # Get last N conversations
                stage = doc.metadata.get("stage", "unknown")
                summary_parts.append(f"[{stage}] {doc.page_content}")
            
            return "\n".join(summary_parts)
            
        except Exception as e:
            print(f"[USER_RAG] Error getting summary for user {user_id}: {e}")
            return ""
    
    def clear_user_history(self, user_id: str):
        """Clear all conversation history for a user"""
        if user_id in self.user_vectorstores:
            try:
                user_persist_dir = os.path.join(self.base_persist_directory, f"user_{user_id}")
                if os.path.exists(user_persist_dir):
                    import shutil
                    shutil.rmtree(user_persist_dir)
                
                del self.user_vectorstores[user_id]
                print(f"[USER_RAG] Cleared history for user {user_id}")
            except Exception as e:
                print(f"[USER_RAG] Error clearing history for user {user_id}: {e}")

class WordDocumentLoader:
    """Custom loader for Word documents (.docx files)"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """Load content from Word document"""
        if not DOCX_AVAILABLE:
            print(f"[WORD_LOADER] python-docx not available, skipping {self.file_path}")
            return []
        
        try:
            doc = DocxDocument(self.file_path)
            
            # Extract all text from paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    full_text.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            content = "\n".join(full_text)
            
            print(f"[WORD_LOADER] Successfully loaded {self.file_path}, content length: {len(content)}")
            
            metadata = {
                "source": self.file_path,
                "type": "word_document"
            }
            
            return [Document(page_content=content, metadata=metadata)]
            
        except Exception as e:
            print(f"[WORD_LOADER] Error loading {self.file_path}: {e}")
            return []
    
    def lazy_load(self):
        """Implement lazy_load for compatibility"""
        docs = self.load()
        for doc in docs:
            yield doc

class CompanyDocumentsRAG:
    """RAG system for company documents and product information"""
    
    def __init__(self, persist_directory: str = "company_docs_vectors"):
        self.persist_directory = persist_directory
        self.gemini_api_key = os.getenv("GEMINI_API_KEY")
        self.groq_api_key = os.getenv("GROQ_API_KEY")
        
        if self.gemini_api_key:
            self.embeddings = GoogleGenerativeAIEmbeddings(
                model="models/embedding-001",
                google_api_key=self.gemini_api_key
            )
        else:
            self.embeddings = None
            print("[COMPANY_RAG] Warning: GEMINI_API_KEY not found. Company docs RAG disabled.")
        
        if self.groq_api_key:
            self.llm = ChatGroq(
                model="llama-3.3-70b-versatile",
                groq_api_key=self.groq_api_key,
                temperature=0.7
            )
        else:
            self.llm = None
            print("[COMPANY_RAG] Warning: GROQ_API_KEY not found. Response generation disabled.")
        
        self.vectorstore = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        
        self._load_or_create_vectorstore()
    
    def _load_or_create_vectorstore(self):
        """Load existing or create new vectorstore for company documents"""
        if not self.embeddings:
            return
        
        try:
            if os.path.exists(self.persist_directory):
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings,
                    collection_name="company_documents"
                )
                print(f"[COMPANY_RAG] Loaded existing company documents vectorstore")
            else:
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings,
                    collection_name="company_documents"
                )
                print(f"[COMPANY_RAG] Created new company documents vectorstore")
        except Exception as e:
            print(f"[COMPANY_RAG] Error with vectorstore: {e}")
            self.vectorstore = None
    
    def load_company_documents(self, docs_directory: str = "src/react_agent/docs"):
        """Load company documents from directory, supporting both .txt and .docx files"""
        if not self.embeddings:
            print("[COMPANY_RAG] Cannot load documents - embeddings not available")
            return
        
        # Ensure vectorstore is available
        if not self.vectorstore:
            try:
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings,
                    collection_name="company_documents"
                )
                print("[COMPANY_RAG] Created new vectorstore for document loading")
            except Exception as e:
                print(f"[COMPANY_RAG] Failed to create vectorstore: {e}")
                return
        
        if not os.path.exists(docs_directory):
            print(f"[COMPANY_RAG] Documents directory {docs_directory} not found")
            return
        
        try:
            all_documents = []
            
            # Process .txt files
            txt_files = []
            for root, dirs, files in os.walk(docs_directory):
                for file in files:
                    if file.endswith('.txt'):
                        txt_files.append(os.path.join(root, file))
            
            if txt_files:
                print(f"[COMPANY_RAG] Found {len(txt_files)} .txt files")
                loader = DirectoryLoader(
                    docs_directory,
                    glob="**/*.txt",
                    show_progress=True
                )
                txt_documents = loader.load()
                all_documents.extend(txt_documents)
                print(f"[COMPANY_RAG] Loaded {len(txt_documents)} txt documents")
            
            # Process .docx files
            docx_files = []
            for root, dirs, files in os.walk(docs_directory):
                for file in files:
                    if file.endswith('.docx'):
                        docx_files.append(os.path.join(root, file))
            
            if docx_files:
                print(f"[COMPANY_RAG] Found {len(docx_files)} .docx files: {docx_files}")
                for docx_file in docx_files:
                    loader = WordDocumentLoader(docx_file)
                    docx_documents = loader.load()
                    all_documents.extend(docx_documents)
                    print(f"[COMPANY_RAG] Loaded Word document: {docx_file}")
            
            print(f"[COMPANY_RAG] Total loaded documents: {len(all_documents)}")
            
            if not all_documents:
                print("[COMPANY_RAG] No documents found")
                return
            
            # Print sample content from Word documents
            for doc in all_documents:
                if doc.metadata.get("type") == "word_document":
                    source = doc.metadata.get("source", "unknown")
                    filename = os.path.basename(source)
                    print(f"[COMPANY_RAG] Word document content preview ({filename}): {doc.page_content[:300]}...")
            
            # Split documents
            texts = self.text_splitter.split_documents(all_documents)
            print(f"[COMPANY_RAG] Split into {len(texts)} chunks")
            
            # Show chunk distribution by source
            chunk_sources = {}
            for chunk in texts:
                source = chunk.metadata.get('source', 'unknown')
                filename = os.path.basename(source)
                if filename not in chunk_sources:
                    chunk_sources[filename] = 0
                chunk_sources[filename] += 1
            
            print(f"[COMPANY_RAG] Chunks per source: {chunk_sources}")
            
            # Clear existing documents and add new ones
            try:
                # Try to get current document count
                current_count = 0
                try:
                    collection = self.vectorstore._collection
                    current_count = collection.count()
                    print(f"[COMPANY_RAG] Current vectorstore has {current_count} documents")
                except:
                    print("[COMPANY_RAG] Could not check current document count")
                
                # If empty or we want to reload, clear and add documents
                if current_count == 0 or len(texts) > current_count:
                    try:
                        self.vectorstore.delete_collection()
                        print("[COMPANY_RAG] Cleared existing vectorstore")
                    except:
                        print("[COMPANY_RAG] Could not clear existing vectorstore")
                    
                    # Recreate vectorstore
                    self.vectorstore = Chroma(
                        persist_directory=self.persist_directory,
                        embedding_function=self.embeddings,
                        collection_name="company_documents"
                    )
                    print("[COMPANY_RAG] Recreated vectorstore")
                    
                    # Add documents
                    self.vectorstore.add_documents(texts)
                    print(f"[COMPANY_RAG] Added {len(texts)} chunks to vectorstore")
                else:
                    print(f"[COMPANY_RAG] Vectorstore already has {current_count} documents, skipping reload")
                
            except Exception as e:
                print(f"[COMPANY_RAG] Error managing vectorstore: {e}")
                # Try to add documents anyway
                try:
                    self.vectorstore.add_documents(texts)
                    print(f"[COMPANY_RAG] Added {len(texts)} chunks to vectorstore (fallback)")
                except Exception as add_error:
                    print(f"[COMPANY_RAG] Failed to add documents: {add_error}")
            
        except Exception as e:
            print(f"[COMPANY_RAG] Error loading documents: {e}")
            import traceback
            traceback.print_exc()
    
    def get_product_suggestions(self, user_query: str, user_interest: str = "", k: int = 5) -> List[Dict]:
        """Get product suggestions based on user query and interests"""
        if not self.vectorstore:
            return []
        
        # Enhance query with user interest
        enhanced_query = f"{user_query} {user_interest}".strip()
        
        try:
            docs = self.vectorstore.similarity_search(enhanced_query, k=k)
            suggestions = []
            
            for doc in docs:
                suggestions.append({
                    "content": doc.page_content,
                    "source": doc.metadata.get("source", "unknown"),
                    "relevance_score": "high"  # Could implement actual scoring
                })
            
            print(f"[COMPANY_RAG] Found {len(suggestions)} product suggestions")
            return suggestions
            
        except Exception as e:
            print(f"[COMPANY_RAG] Error getting suggestions: {e}")
            return []
    
    def generate_personalized_recommendation(self, user_data: Dict, suggestions: List[Dict]) -> str:
        """Generate personalized recommendation using LLM"""
        if not self.llm or not suggestions:
            return "I don't have enough information to make personalized recommendations right now."
        
        # Prepare context
        context = "\n".join([s["content"] for s in suggestions[:3]])  # Top 3 suggestions
        
        user_profile = f"""
        Name: {user_data.get('name', 'User')}
        Age: {user_data.get('age', 'Not specified')}
        Country: {user_data.get('country', 'Not specified')}
        Interest: {user_data.get('interest', 'Not specified')}
        """
        
        prompt = f"""Based on the user profile and company information, provide personalized product recommendations.

User Profile:
{user_profile}

Company Information:
{context}

Please provide 2-3 specific product recommendations with brief explanations of why they suit this user."""
        
        try:
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            print(f"[COMPANY_RAG] Error generating recommendation: {e}")
            return "I'm having trouble generating recommendations right now. Please try again later."

class DualRAGSystem:
    """Unified system managing both user chat history and company documents"""
    
    def __init__(self):
        self.chat_history_rag = UserChatHistoryRAG()
        self.company_docs_rag = CompanyDocumentsRAG()
        
        print("[DUAL_RAG] Initialized dual RAG system")
        
        # Load company documents on initialization with retries
        self._initialize_company_documents()
    
    def _initialize_company_documents(self):
        """Initialize company documents with error handling"""
        max_retries = 3
        for attempt in range(max_retries):
            try:
                print(f"[DUAL_RAG] Attempting to load company documents (attempt {attempt + 1}/{max_retries})")
                self.company_docs_rag.load_company_documents()
                
                # Test if documents were loaded
                test_suggestions = self.company_docs_rag.get_product_suggestions("test", k=1)
                if test_suggestions:
                    print(f"[DUAL_RAG] ✅ Company documents loaded successfully")
                    return
                else:
                    print(f"[DUAL_RAG] ⚠️ Documents loaded but no content found")
                    if attempt < max_retries - 1:
                        # Try to recreate vectorstore
                        try:
                            if self.company_docs_rag.vectorstore:
                                self.company_docs_rag.vectorstore.delete_collection()
                        except:
                            pass
                        
                        # Recreate vectorstore
                        if self.company_docs_rag.embeddings:
                            from langchain_community.vectorstores import Chroma
                            self.company_docs_rag.vectorstore = Chroma(
                                persist_directory=self.company_docs_rag.persist_directory,
                                embedding_function=self.company_docs_rag.embeddings,
                                collection_name="company_documents"
                            )
                        continue
                    
            except Exception as e:
                print(f"[DUAL_RAG] Error loading documents (attempt {attempt + 1}): {e}")
                if attempt < max_retries - 1:
                    continue
                else:
                    print(f"[DUAL_RAG] ❌ Failed to load company documents after {max_retries} attempts")
                    break
    
    def store_user_conversation(self, user_id: str, user_message: str, agent_response: str, 
                               stage: str = "unknown", metadata: Dict = None):
        """Store conversation in user's chat history"""
        self.chat_history_rag.store_conversation(user_id, user_message, agent_response, stage, metadata)
    
    def get_user_context(self, user_id: str, current_query: str, k: int = 5) -> str:
        """Get relevant user context from chat history"""
        history = self.chat_history_rag.get_relevant_history(user_id, current_query, k)
        if not history:
            return ""
        
        context_parts = []
        for h in history:
            context_parts.append(f"Previous conversation: {h['content']}")
        
        return "\n".join(context_parts)
    
    def get_product_recommendations(self, user_data: Dict, query: str = "product recommendation") -> str:
        """Get personalized product recommendations"""
        user_interest = user_data.get('interest', '')
        suggestions = self.company_docs_rag.get_product_suggestions(query, user_interest)
        
        if suggestions:
            return self.company_docs_rag.generate_personalized_recommendation(user_data, suggestions)
        else:
            return "I don't have specific product information available right now."
    
    def clear_user_data(self, user_id: str):
        """Clear all user-specific data"""
        self.chat_history_rag.clear_user_history(user_id)
    
    def get_system_status(self) -> Dict:
        """Get status of both RAG systems"""
        return {
            "chat_history_rag": {
                "embeddings_available": self.chat_history_rag.embeddings is not None,
                "active_users": len(self.chat_history_rag.user_vectorstores)
            },
            "company_docs_rag": {
                "embeddings_available": self.company_docs_rag.embeddings is not None,
                "llm_available": self.company_docs_rag.llm is not None,
                "vectorstore_available": self.company_docs_rag.vectorstore is not None
            }
        }

# Global instance
dual_rag_system = DualRAGSystem() 